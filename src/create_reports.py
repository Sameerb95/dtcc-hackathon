import json
import re
import os
from pathlib import Path # Import Path
from langchain.prompts.chat import ChatPromptTemplate
from langchain.prompts.chat import HumanMessagePromptTemplate
from docx import Document
# Assuming VertexAILangchainLLM is in a place accessible by this script
try:
    from llm_call import VertexAILangchainLLM
except ImportError:
    try:
        from .llm_call import VertexAILangchainLLM
    except ImportError:
        print("Warning: Could not import VertexAILangchainLLM. LLM calls will fail.")
        VertexAILangchainLLM = None

# Assuming get_report_prompt is defined in get_reporting_prompt.py
try:
    from get_reporting_prompt import get_report_prompt
except ImportError:
    try:
        from .get_reporting_prompt import get_report_prompt
    except ImportError:
        print("Warning: Could not import get_report_prompt. KOP from diff will use a default.")
        def get_report_prompt(): # Fallback basic prompt
            return "Generate a KOP document from the following JSON data: {difference_json}"


def get_kop_generation_prompt_for_approved_items(): # Renamed for clarity
    """
    Prompt for generating KOP from a list of approved items.
    """
    prompt = """**Objective:** You are tasked with creating a Key Operating Procedures (KOP) document based on a list of approved changes. Each approved change includes a title, confidence level, a description of the change, and recommended next steps.

            Your KOP document should be structured for clarity and actionability by operational personnel. Please ensure it includes:

            1.  **Overall Introduction:** Briefly summarize the purpose of this KOP document, highlighting that it's based on recently approved guidelines or analysis changes.
            2.  **Procedure Sections:** For each approved item provided, create a distinct section in the KOP. Each section should clearly detail:
                *   **Change Title/Reference:** Use the title of the approved item.
                *   **Summary of Approved Change:** Briefly explain what the change entails, based on its description.
                *   **Operational Steps/Procedures:** Convert the 'Recommended Next Steps' into clear, actionable, step-by-step instructions for operational teams. If next steps are general, infer specific operational actions where possible.
                *   **Confidence Level (Optional Information):** You may note the confidence level associated with the change if relevant for prioritization or attention.
            3.  **Formatting:** Use clear headings for each section. Employ bullet points or numbered lists for step-by-step procedures to enhance readability.
            4.  **Language:** Use precise, unambiguous language suitable for an operational guide.

            Input Data
            ----------------------
            {{approved_items_data_string}}

            Based on the approved changes detailed above, generate the Key Operating Procedures (KOP) document.
            Focus on translating the 'Next Steps' into concrete operational actions.
            """
    return prompt

def markdown_to_docx(doc , text: str):
    """
    Appends markdown text to an existing docx.Document object.
    Handles headings (H1-H3), lists (bullet and numbered), and bold text.
    """
    lines = text.split('\n')
    current_paragraph = None

    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("#").strip(), level=3)
            current_paragraph = None 
            continue
        elif stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("#").strip(), level=2)
            current_paragraph = None
            continue
        elif stripped.startswith("#"):
            doc.add_heading(stripped.lstrip("#").strip(), level=1)
            current_paragraph = None
            continue
        
        is_bullet = stripped.startswith("- ") or stripped.startswith("* ")
        is_numbered = re.match(r'^\d+\.\s', stripped)

        if is_bullet:
            doc.add_paragraph(stripped[2:], style='ListBullet')
            current_paragraph = None
            continue
        elif is_numbered:
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='ListNumber')
            current_paragraph = None
            continue
        
        if not stripped: # Empty line usually means a paragraph break in markdown
            current_paragraph = None # Reset current paragraph
            # doc.add_paragraph() # Optionally add an empty paragraph for spacing, or let next text create it
            continue

        if current_paragraph is None:
             current_paragraph = doc.add_paragraph()
        else: # If it's a continuation of text on a new line but not a special markdown element
            current_paragraph.add_run("\n") # Treat as a soft break if desired, or append to existing run

        parts = re.split(r'(\*\*.*?\*\*)', stripped) 
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                if part[2:-2]: # Ensure there's text between **
                    bold_run = current_paragraph.add_run(part[2:-2])
                    bold_run.bold = True
            elif part: 
                current_paragraph.add_run(part)
    return doc


def send_diff_to_llm(graph_difference_json_string, output_directory=".", filename="KOP_AWPR_from_Diff.docx"):
    """
    Sends graph difference JSON to LLM, generates KOP, and saves it to a DOCX file.
    Returns the full path to the saved document or None on failure.
    """
    if not VertexAILangchainLLM:
        print("Error: VertexAILangchainLLM is not available. Cannot process diff.")
        return None
 
    print(f"--- Preparing KOP from diff JSON ---")
    messages = []
    template = get_report_prompt()
    human_template = HumanMessagePromptTemplate.from_template(template)
    messages.append(human_template) 
    chat_prompt = ChatPromptTemplate.from_messages(messages)
    prompt_messages = chat_prompt.format_prompt(difference_json=graph_difference_json_string).to_messages()
    request_content_str = "".join([msg.content for msg in prompt_messages])

    print(f"Formatted Prompt for LLM (KOP from Diff): {request_content_str[:500]}...")

    llm = VertexAILangchainLLM({})
    try:
        response = llm._call(prompt=request_content_str)
        print("\n--- LLM Response (KOP from Diff) ---")
        print(response)

        doc = Document()
        doc.add_heading("Key Operating Procedure (KOP) - Based on Document Differences", level=0)
        markdown_to_docx(doc, response) # Use the refined markdown_to_docx
        
        # Ensure the KOP subfolder exists within the output_directory
        kop_output_folder = Path(output_directory) / "KOP"
        kop_output_folder.mkdir(parents=True, exist_ok=True)
        
        full_path = kop_output_folder / filename

        try:
            doc.save(str(full_path)) # doc.save expects a string path
            print(f"KOP (from diff) saved successfully to: {full_path}")
            return str(full_path)
        except Exception as e:
            print(f"Error saving KOP (from diff) document: {e}")
            return None

    except Exception as e:
        print(f"An error occurred while calling LLM for KOP (from diff): {e}")
        return None

def generate_kop_from_approved_items(approved_items_list, output_directory=".", filename="KOP_From_Approved_Items.docx"):
    """
    Generates a KOP document (DOCX) from a list of approved items.
    Returns the full path to the saved document or None on failure.
    """
    if not VertexAILangchainLLM:
        print("Error: VertexAILangchainLLM is not available. Cannot generate KOP.")
        return None

    if not approved_items_list:
        print("No approved items provided. KOP document will not be generated.")
        return None

    print(f"--- Preparing KOP from {len(approved_items_list)} approved items ---")

    approved_items_data_str = ""
    for i, item in enumerate(approved_items_list):
        approved_items_data_str += f"Item {i+1}:\n"
        approved_items_data_str += f"  Title: {item.get('title_display', item.get('title_text', 'N/A'))}\n"
        approved_items_data_str += f"  Confidence: {item.get('confidence', 'N/A')}\n"
        approved_items_data_str += f"  Description: {item.get('description', 'N/A')}\n"
        approved_items_data_str += f"  Recommended Next Steps: {item.get('next_steps', 'N/A')}\n\n"

    print(approved_items_data_str)

    messages = []
    template = get_report_prompt()
    human_template = HumanMessagePromptTemplate.from_template(template)
    messages.append(human_template)
    chat_prompt = ChatPromptTemplate.from_messages(messages)

    prompt_messages = chat_prompt.format_prompt(approved_items_data_string=approved_items_data_str.strip()).to_messages()
    request_content_str = "".join([msg.content for msg in prompt_messages])

    print(f"Formatted Prompt for KOP (Approved Items) LLM: {request_content_str[:1000]}...")

    llm = VertexAILangchainLLM({})
    try:
        llm_response_markdown = llm._call(prompt=request_content_str)
        print("\n--- LLM Response for KOP (Approved Items) ---")
        print(llm_response_markdown)

        doc = Document()
        doc.add_heading("Key Operating Procedures (KOP) - Based on Approved Analysis Items", level=0)
        markdown_to_docx(doc, llm_response_markdown) # Use the refined markdown_to_docx
        
        # Ensure the KOP subfolder exists
        kop_output_folder = Path(output_directory) / "KOP"
        kop_output_folder.mkdir(parents=True, exist_ok=True)
        
        full_path = kop_output_folder / filename

        try:
            doc.save(str(full_path)) # doc.save expects a string path
            print(f"KOP (Approved Items) Document saved successfully to: {full_path}")
            return str(full_path)
        except Exception as e:
            print(f"Error saving KOP (Approved Items) document: {e}")
            return None

    except Exception as e:
        print(f"An error occurred while calling LLM for KOP (Approved Items) generation: {e}")
        return None


# if __name__ == "__main__":
#     # Path to your merged JSON file (for testing send_diff_to_llm)
#     merged_json_path = "merged_graph_output.json" # Adjust if necessary
#     test_output_directory = "generated_reports_test" # For testing

#     if os.path.exists(merged_json_path):
#         with open(merged_json_path, 'r') as f:
#             graph_diff_json_content_string = f.read()

#             if graph_diff_json_content_string:
#                 print("\n--- Testing KOP generation from DIFF JSON ---")
#                 saved_path_diff = send_diff_to_llm(graph_diff_json_content_string, output_directory=test_output_directory)
#                 if saved_path_diff:
#                     print(f"KOP from diff saved to: {saved_path_diff}")
#                 else:
#                     print("KOP generation from diff failed.")
#             else:
#                 print(f"Error: The file '{merged_json_path}' is empty.")
#     else:
#         print(f"Warning: Merged JSON file not found at '{merged_json_path}'. Skipping KOP from diff test.")

#     # Example test for the generate_kop_from_approved_items function
#     print("\n--- Testing KOP generation from APPROVED ITEMS ---")
#     dummy_approved_items = [
#         {
#             "title_display": "(1) Test Item Alpha (High Confidence)",
#             "title_text": "Test Item Alpha",
#             "confidence": "High Confidence",
#             "description": "This is a test description for item Alpha, detailing its changes.",
#             "next_steps": "1. Implement new protocol. 2. Train staff. 3. Monitor results."
#         },
#         {
#             "title_display": "(2) Test Item Beta (Medium Confidence)",
#             "title_text": "Test Item Beta",
#             "confidence": "Medium Confidence",
#             "description": "Item Beta involves updating the user interface.",
#             "next_steps": "1. Design mockups. 2. Get user feedback. 3. Develop and deploy."
#         }
#     ]
#     saved_path_approved = generate_kop_from_approved_items(
#         dummy_approved_items, 
#         output_directory=test_output_directory,
#         filename="Test_KOP_From_Approved.docx"
#     )
#     if saved_path_approved:
#         print(f"KOP from approved items saved to: {saved_path_approved}")
#     else:
#         print("KOP generation from approved items failed.")
